# ingestion/document_loader.py
# ─────────────────────────────────────────────────────────────────────────────
# Stage 1 — Document Ingestion
#
# Responsibilities:
#   • Detect file type (PDF native / PDF scanned / image / DOCX)
#   • Route to the correct parser
#   • Return a unified list of PageData objects (page_num, text, tables, source)
#   • Auto-detect document language for optimal OCR
#
# Supported formats:
#   PDF (native text) → pdfplumber
#   PDF (scanned)     → PaddleOCR PP‑OCRv4
#   Images            → Vision LLM (llava) with PaddleOCR fallback
#   DOCX              → python-docx
#   TXT / MD          → plain read
#   Web URLs          → BeautifulSoup
# ─────────────────────────────────────────────────────────────────────────────

from __future__ import annotations

import io
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

import pdfplumber
import fitz                          # PyMuPDF
from docx import Document as DocxDoc
from loguru import logger

from config import (
    OCR_LANGUAGE, OCR_USE_ANGLE_CLS, ALLOWED_EXTENSIONS,
    MAX_FILE_SIZE_MB, MULTILINGUAL_MODE,
)


# ── Data container returned by every parser ──────────────────────────────────

@dataclass
class PageData:
    """Unified output from any document parser."""
    source: str                      # file name
    page_num: int                    # 1-based
    text: str                        # extracted plain text
    tables: List[List[List[str]]] = field(default_factory=list)
    method: str = "native"           # "native" | "ocr" | "docx" | "text" | "vision" | "web"
    language: str = "unknown"        # auto-detected language code


# ── Language Detection ───────────────────────────────────────────────────────

def _detect_language(text: str) -> str:
    """Detect the language of a text snippet. Returns ISO 639-1 code."""
    if not text or len(text.strip()) < 20:
        return "unknown"
    try:
        from langdetect import detect
        return detect(text[:1000])  # Only check first 1000 chars for speed
    except Exception:
        return "unknown"


def _get_ocr_lang(detected_lang: str) -> str:
    """Map detected language code to PaddleOCR language code."""
    lang_map = {
        "en": "en", "hi": "hi", "mr": "mr", "ta": "ta", "te": "te",
        "zh-cn": "ch", "zh-tw": "ch", "ja": "japan", "ko": "korean",
        "fr": "fr", "de": "german", "es": "es", "pt": "pt",
        "ar": "ar", "ru": "ru", "it": "it",
    }
    return lang_map.get(detected_lang, OCR_LANGUAGE)


# ── Lazy-load PaddleOCR so startup is fast when OCR is not needed ────────────

_ocr_engines = {}


def _get_ocr(lang: str = None):
    """Load PaddleOCR engine lazily (downloads ~45 MB on first run)."""
    lang = lang or OCR_LANGUAGE
    if lang not in _ocr_engines:
        from paddleocr import PaddleOCR
        logger.info(f"Loading PaddleOCR engine for '{lang}' (first run downloads ~45 MB)…")
        _ocr_engines[lang] = PaddleOCR(
            use_angle_cls=OCR_USE_ANGLE_CLS,
            lang=lang,
            show_log=False,
        )
        logger.info(f"PaddleOCR ({lang}) ready.")
    return _ocr_engines[lang]


# ── Helper: is a PDF page text-selectable? ───────────────────────────────────

def _page_has_native_text(page: fitz.Page, min_chars: int = 20) -> bool:
    """Returns True if the page already has enough selectable text."""
    return len(page.get_text("text").strip()) >= min_chars


# ── Parser: native (text-selectable) PDF ─────────────────────────────────────

def _parse_native_pdf(path: Path) -> List[PageData]:
    """Extract text and tables from a text-selectable PDF using pdfplumber."""
    pages: List[PageData] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text   = page.extract_text() or ""
            tables = page.extract_tables() or []
            clean_tables = [
                [[cell or "" for cell in row] for row in tbl]
                for tbl in tables
            ]
            lang = _detect_language(text) if MULTILINGUAL_MODE else "en"
            pages.append(PageData(
                source=path.name, page_num=i,
                text=text, tables=clean_tables, method="native",
                language=lang,
            ))
    logger.debug(f"Native PDF parsed: {len(pages)} pages from {path.name}")
    return pages


# ── Parser: scanned PDF / image  (OCR path) ──────────────────────────────────

def _parse_scanned_pdf(path: Path) -> List[PageData]:
    """Render each PDF page as an image and run PaddleOCR on it."""
    doc = fitz.open(str(path))
    pages: List[PageData] = []

    # Try to detect language from any native text first
    ocr_lang = OCR_LANGUAGE
    if MULTILINGUAL_MODE:
        for page in doc:
            sample_text = page.get_text("text").strip()
            if len(sample_text) > 20:
                detected = _detect_language(sample_text)
                ocr_lang = _get_ocr_lang(detected)
                break
        doc.close()
        doc = fitz.open(str(path))

    ocr = _get_ocr(ocr_lang)

    for i, page in enumerate(doc, start=1):
        # Render at 2× resolution for better OCR accuracy
        mat  = fitz.Matrix(2.0, 2.0)
        pix  = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
        img_bytes = pix.tobytes("png")

        result = ocr.ocr(img_bytes, cls=True)
        lines  = []
        if result and result[0]:
            for line in result[0]:
                text_str, conf = line[1]
                if conf > 0.5:
                    lines.append(text_str)

        text = "\n".join(lines)
        lang = _detect_language(text) if MULTILINGUAL_MODE else ocr_lang
        pages.append(PageData(
            source=path.name, page_num=i,
            text=text, tables=[], method="ocr",
            language=lang,
        ))

    doc.close()
    logger.debug(f"Scanned PDF parsed (OCR): {len(pages)} pages from {path.name}")
    return pages


# ── Parser: image file (.png / .jpg / .jpeg / .tiff) ─────────────────────────

def _parse_image(path: Path) -> List[PageData]:
    """Run PaddleOCR directly on an image file."""
    ocr    = _get_ocr()
    result = ocr.ocr(str(path), cls=True)
    lines  = []
    if result and result[0]:
        for line in result[0]:
            text_str, conf = line[1]
            if conf > 0.5:
                lines.append(text_str)

    text = "\n".join(lines)
    lang = _detect_language(text) if MULTILINGUAL_MODE else "en"
    return [PageData(
        source=path.name, page_num=1,
        text=text, tables=[], method="ocr",
        language=lang,
    )]


def _parse_image_vision(path: Path) -> List[PageData]:
    """Parse image using Vision LLM (like llava) via Ollama, with fallback to OCR."""
    import base64
    from langchain_ollama import ChatOllama
    from langchain_core.messages import HumanMessage
    
    with open(path, "rb") as f:
        img_b64 = base64.b64encode(f.read()).decode("utf-8")
        
    try:
        logger.info(f"Attempting Vision LLM (Llava) extraction for {path.name}")
        llm = ChatOllama(model="llava", temperature=0)
        msg = HumanMessage(
            content=[
                {"type": "text", "text": "Extract all text, data, tables, and describe any charts/graphs in this image in detail. Reply ONLY with the extracted content and description."},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}}
            ]
        )
        response = llm.invoke([msg])
        text = response.content.strip() if hasattr(response, "content") else str(response)
        lang = _detect_language(text) if MULTILINGUAL_MODE else "en"
        
        return [PageData(
            source=path.name, page_num=1,
            text=text, tables=[], method="vision",
            language=lang,
        )]
    except Exception as e:
        logger.warning(f"Vision LLM failed ({e}), falling back to standard OCR for {path.name}")
        return _parse_image(path)


# ── Parser: DOCX ─────────────────────────────────────────────────────────────

def _parse_docx(path: Path) -> List[PageData]:
    """Extract paragraphs and tables from a Word document."""
    doc   = DocxDoc(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]

    # Extract tables from DOCX
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)

    text = "\n".join(paras)
    lang = _detect_language(text) if MULTILINGUAL_MODE else "en"
    return [PageData(
        source=path.name, page_num=1,
        text=text, tables=tables, method="docx",
        language=lang,
    )]


# ── Parser: plain text / markdown ─────────────────────────────────────────────

def _parse_text(path: Path) -> List[PageData]:
    """Read a plain text or markdown file."""
    text = path.read_text(encoding="utf-8", errors="replace")
    lang = _detect_language(text) if MULTILINGUAL_MODE else "en"
    return [PageData(
        source=path.name, page_num=1,
        text=text, tables=[], method="text",
        language=lang,
    )]


# ── Parser: Web URL ──────────────────────────────────────────────────────────

def parse_url(url: str) -> List[PageData]:
    """Fetch a web page and extract the text content using BeautifulSoup."""
    try:
        import requests
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("Web crawling requires 'requests' and 'beautifulsoup4'. Run: pip install requests beautifulsoup4")

    logger.info(f"Fetching URL: {url}")
    headers = {"User-Agent": "Mozilla/5.0 Document AI RAG Pipeline"}
    resp = requests.get(url, headers=headers, timeout=10)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    
    # Remove javascript, styles, and empty tags
    for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
        script.decompose()
        
    text = soup.get_text(separator="\n\n")
    
    # Clean up empty lines
    lines = (line.strip() for line in text.splitlines())
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    cleaned_text = "\n".join(chunk for chunk in chunks if chunk)
    
    # Use the domain + title as the source name
    from urllib.parse import urlparse
    domain = urlparse(url).netloc
    title = soup.title.string.strip() if soup.title else domain
    source_name = f"{domain} - {title}"
    lang = _detect_language(cleaned_text) if MULTILINGUAL_MODE else "en"

    return [PageData(
        source=source_name, page_num=1,
        text=cleaned_text, tables=[], method="web",
        language=lang,
    )]


# ── Public API ────────────────────────────────────────────────────────────────

def load_document(path: str | Path) -> List[PageData]:
    """
    Auto-detect file type and return a list of PageData objects.

    Supports:  PDF, PNG, JPG, JPEG, TIFF, BMP, WEBP, DOCX, TXT, MD

    Usage:
        pages = load_document("invoice.pdf")
        for page in pages:
            print(page.page_num, page.text[:100])

    Raises:
        FileNotFoundError: if the file does not exist
        ValueError: if the file type is unsupported or exceeds size limit
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    # File size check
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(
            f"File too large: {size_mb:.1f} MB  "
            f"(max allowed: {MAX_FILE_SIZE_MB} MB)"
        )

    suffix = path.suffix.lower()
    logger.info(f"Loading document: {path.name}  (type={suffix}, size={size_mb:.1f} MB)")
    t0 = time.perf_counter()

    if suffix == ".pdf":
        # Quick check: open with PyMuPDF and sample first 3 pages
        doc = fitz.open(str(path))
        sample_pages = min(3, len(doc))
        native_count = sum(
            1 for i in range(sample_pages)
            if _page_has_native_text(doc[i])
        )
        doc.close()

        if native_count == sample_pages:
            logger.info("→ Native PDF detected, using pdfplumber parser")
            pages = _parse_native_pdf(path)
        else:
            logger.info("→ Scanned / image PDF detected, using OCR parser")
            pages = _parse_scanned_pdf(path)

    elif suffix in {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}:
        logger.info("→ Image file, attempting Vision LLM with OCR fallback")
        pages = _parse_image_vision(path)

    elif suffix in {".docx", ".doc"}:
        logger.info("→ DOCX file, using python-docx parser")
        pages = _parse_docx(path)

    elif suffix in {".txt", ".md"}:
        logger.info("→ Text file, reading directly")
        pages = _parse_text(path)

    else:
        raise ValueError(
            f"Unsupported file type: {suffix}  "
            f"(supported: pdf, png, jpg, jpeg, tiff, bmp, webp, docx, txt, md)"
        )

    elapsed = time.perf_counter() - t0
    total_chars = sum(len(p.text) for p in pages)
    total_tables = sum(len(p.tables) for p in pages)
    detected_langs = {p.language for p in pages if p.language != "unknown"}
    logger.info(
        f"Document loaded in {elapsed:.2f}s — "
        f"{len(pages)} pages, {total_chars:,} chars, {total_tables} tables"
        f"{f', languages: {detected_langs}' if detected_langs else ''}"
    )
    return pages
